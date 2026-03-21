import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_cell_margins(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin in ['top', 'left', 'bottom', 'right']:
        if margin in kwargs:
            node = OxmlElement(f'w:{margin}')
            node.set(qn('w:w'), str(kwargs[margin]))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)

def add_blue_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.size = Pt(14)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '003366')
    pBdr.append(bottom)
    pPr.append(pBdr)

def generate_docx_report(user_name: str, test_title: str, score: int, answers_list: list, category_scores: dict = None, report_config: dict = None) -> io.BytesIO:
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    
    # Header & Footer
    sec = doc.sections[0]
    sec.header.paragraphs[0].text = "ProfDNK | Система профориентационного анализа"
    p_footer = sec.footer.paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = p_footer.add_run()
    run_footer.add_text("Страница ")
    fld1 = OxmlElement('w:fldChar')
    fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')
    run_footer._r.append(fld1)
    run_footer._r.append(instr)
    run_footer._r.append(fld2)
    
    # Logo Placeholder
    t_logo = doc.add_table(rows=1, cols=1)
    t_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_logo = t_logo.cell(0, 0)
    p_c_logo = c_logo.paragraphs[0]
    p_c_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_c_logo = p_c_logo.add_run("[ LOGO ПрофДНК ]")
    run_c_logo.font.size = Pt(16)
    run_c_logo.font.color.rgb = RGBColor(0, 51, 102)
    run_c_logo.bold = True
    tbl_logo_pr = t_logo._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for bdr in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{bdr}')
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), '12') # 1.5pt
        node.set(qn('w:space'), '0')
        node.set(qn('w:color'), '003366')
        tblBorders.append(node)
    tbl_logo_pr.append(tblBorders)
    set_cell_margins(c_logo, top=144, left=144, bottom=144, right=144)
    
    doc.add_paragraph()
    
    # Шапка
    heading = doc.add_heading(level=0)
    run = heading.add_run("Результаты профориентационного тестирования ПрофДНК")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    import datetime
    current_date = datetime.datetime.now().strftime("%d.%m.%Y %H:%M")
    
    # Личные данные
    add_blue_heading(doc, "Анализируемый профиль")
    doc.add_paragraph(f"ФИО: {user_name}")
    doc.add_paragraph(f"Дата прохождения: {current_date}")
    doc.add_paragraph(f"Тест: {test_title}")
    
    p_score = doc.add_paragraph()
    p_score.add_run("Итоговое соответствие: ").bold = True
    p_score.add_run(f"{score}%")
    
    if report_config is None:
        report_config = {"show_table": True, "show_chart": False, "show_interpretation": True}

    if report_config.get("show_table", True):
        if category_scores and len(category_scores) > 0:
            doc.add_paragraph()
            add_blue_heading(doc, "Профиль компетенций")
            
            cat_table = doc.add_table(rows=1, cols=2)
            cat_table.style = 'Table Grid'
            hdr_cells = cat_table.rows[0].cells
            hdr_cells[0].text = 'Категория / Секция'
            hdr_cells[1].text = 'Уровень (%)'
            
            for cell in hdr_cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'EAEAEA')
                tcPr.append(shd)
                
            for cat, cat_score in category_scores.items():
                row_cells = cat_table.add_row().cells
                row_cells[0].text = str(cat)
                row_cells[1].text = f"{cat_score}%"
                
        doc.add_paragraph()
        add_blue_heading(doc, "Ответы пользователя")
    
    # Таблица
    table = doc.add_table(rows=1, cols=2)
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders2 = OxmlElement('w:tblBorders')
    for bdr in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        node = OxmlElement(f'w:{bdr}')
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), '12') # 1.5pt
        node.set(qn('w:space'), '0')
        node.set(qn('w:color'), 'auto')
        tblBorders2.append(node)
    tblPr.append(tblBorders2)
    
    tblCellMar = OxmlElement('w:tblCellMar')
    for margin in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), "144") # 144 twips = 0.1 inch padding
        node.set(qn('w:type'), "dxa")
        tblCellMar.append(node)
    tblPr.append(tblCellMar)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Вопрос'
    hdr_cells[1].text = 'Вариант ответа'
    
    # Заливка заголовков светло-серая
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'EAEAEA')
        tcPr.append(shd)
    
        for ans in answers_list:
            row_cells = table.add_row().cells
            row_cells[0].text = str(ans.get('question', ''))
            row_cells[1].text = str(ans.get('answer', ''))
        
    
    # Интерпретация
    if report_config.get("show_interpretation", True):
        doc.add_paragraph()
        add_blue_heading(doc, "Краткое заключение")
    
    if score > 70:
        interpretation = "Высокий уровень соответствия профилю. Кандидат обладает выраженными склонностями и компетенциями в данной области."
        color = RGBColor(0, 128, 0)
    elif score >= 40:
        interpretation = "Средний уровень. Требуется дополнительное погружение, но базовые задатки присутствуют."
        color = RGBColor(200, 150, 0)
    else:
        interpretation = "По данному тесту выявлен низкий уровень соответствия. Рекомендуется консультация специалиста для подбора альтернативных направлений."
        color = RGBColor(255, 0, 0)
        
    p_interp = doc.add_paragraph()
    run_interp = p_interp.add_run(interpretation)
    if color:
        run_interp.font.color.rgb = color
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer
